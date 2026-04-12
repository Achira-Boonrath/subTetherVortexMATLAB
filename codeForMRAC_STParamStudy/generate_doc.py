import sys
import subprocess
import os

# Ensure python-docx is installed securely and quickly
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Unified Optimal Control Implementation Details and Theoretical Background', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        "The MATLAB function optControl_DirectUnify.m solves nonlinear optimal control problems using "
        "two different direct transcription methods: Direct Single Shooting and Hermite-Simpson Direct Collocation. "
        "It supports switching between two dynamical models: a Simple Inverted Pendulum and a Cart-Pole System. "
        "By framing trajectory generation as an optimization problem (NLP), it leverages MATLAB's fmincon SQP solver "
        "to find the time-history of control inputs that minimize the required effort and reach the target states."
    )
    
    doc.add_heading('2. Systems Overview and Dynamics', level=1)
    
    doc.add_heading('2.1 Simple Inverted Pendulum', level=2)
    doc.add_paragraph(
        "The simple inverted pendulum consists of a point mass affixed to a massless rod, rotating around a fixed pivot. "
        "The goal is to swing the pendulum from a downward-hanging stable equilibrium to an inverted unstable equilibrium.\n\n"
        "State Variables: Angle (theta) and Angular Velocity (theta_dot).\n"
        "Control Variable: Actuator torque applied at the pivot (u).\n\n"
        "Equations of Motion:"
    )
    doc.add_paragraph("I * theta_ddot = u - m * g * l * sin(theta)", style='Intense Quote')
    doc.add_paragraph(
        "where m is the mass, l is the pole length, g is gravity, and I = m * l^2 is the moment of inertia."
    )
    
    doc.add_heading('2.2 Cart-Pole System', level=2)
    doc.add_paragraph(
        "The cart-pole system places the inverted pendulum on a horizontally translating cart. Because the pivot translates, "
        "the system behaves differently and has unactuated degrees of freedom (the pole itself has no direct torque applied, "
        "it is controlled indirectly via the cart momentum).\n\n"
        "State Variables: Cart Position (x), Cart Velocity (x_dot), Pole Angle (theta), and Pole Angular Velocity (theta_dot).\n"
        "Control Variable: Horizontal linear force applied to the cart (u).\n\n"
        "Equations of Motion (derived via Lagrangian dynamics):"
    )
    doc.add_paragraph(
        "(M + m) * x_ddot + m * l * cos(theta) * theta_ddot = u + m * l * (theta_dot^2) * sin(theta)\n"
        "m * l * cos(theta) * x_ddot + I * theta_ddot = -m * g * l * sin(theta)",
        style='Intense Quote'
    )
    doc.add_paragraph(
        "By organizing these equations into a linear matrix representation (A * x = B), "
        "accelerations are extracted via matrix inversion (MATLAB's backslash operator) during each integration step."
    )
    
    doc.add_heading('3. Optimal Control Transcription Methods', level=1)
    
    doc.add_heading('3.1 Direct Single Shooting', level=2)
    doc.add_paragraph(
        "Direct Single Shooting is a numerical transcription method that parameterizes only the control sequence. "
        "The states are considered dependent variables that evolve implicitly.\n\n"
        "• Decision Variables: Only the piecewise-constant control inputs U at each interval form the optimization vector.\n"
        "• Propagation: To compute the cost and terminal constraints, the system is simulated forward using an ODE solver (such as ode45) from the initial state.\n"
        "• Trade-offs: Because the NLP decision vector is much smaller, fmincon iterations are fast. However, "
        "simulating unstable systems inherently suffers from the \"tail-wagging-the-dog\" effect—tiny changes in early "
        "control inputs produce massive exponential divergence in final states, creating jagged, highly sensitive gradients "
        "that often derail NLP convergence for long time horizons."
    )
    
    doc.add_heading('3.2 Hermite-Simpson Direct Collocation', level=2)
    doc.add_paragraph(
        "Direct Collocation transforms both the control inputs and the state trajectories into explicit independent decision variables.\n\n"
        "• Decision Variables: The optimizer evaluates a massive joint vector bounding specific controls U, node states X, and interval midpoints Xm.\n"
        "• NLP Constraints (Defects): Instead of simulating forward, the optimization uses algebraic constraints called \"defects\" "
        "to enforce the equations of motion across polynomials. It employs 3rd-order accurate Hermite-Simpson approximation:\n\n"
        "   Midpoint Consistency: Evaluates state spline consistency vs numerical midpoints.\n"
        "   Simpson Defect: Enforces numerical integration across [k, k+1] mapped against differences in state endpoints.\n\n"
        "• Trade-offs: Direct Collocation scales effectively into non-linear, rigid, or highly unstable systems (like the Cart-Pole) "
        "because all nodes can update independently without divergent forward integrations. Despite multiplying the number of variables "
        "substantially, modern sparse structure algorithms easily navigate this formulation to reach stable global minimums faster."
    )
    
    doc.add_heading('4. Implementation Details in optControl_DirectUnify.m', level=1)
    doc.add_paragraph(
        "1. Dynamic Bounds Assembly: Depending on the chosen system (pendulum or cartpole), the state tracking size (nx) scales. "
        "Matrices like Qf (Terminal Cost) dynamically map 2 or 4 diagonal entries to penalize target deviations appropriately.\n"
        "2. Function Nesting: Both methods use separate localized closures (objFun_shooting, nonlcon_shooting, objFun_collocation, "
        "nonlcon_collocation). This ensures memory scopes and dimensional layouts inside fmincon correctly match the selected active method configuration.\n"
        "3. Generalized Cost Setup: The tracking objective is split into two components: Running Cost J_run (penalizes high energy U forces) "
        "and Terminal Cost J_final (penalizes mapping discrepancies using the wrapToPi helper avoiding modulus errors)."
    )
    
    # Save the document inside the current working directory
    file_path = os.path.join(os.getcwd(), 'optControl_Documentation.docx')
    doc.save(file_path)
    print(f"Documentation file saved successfully as {file_path}")

if __name__ == "__main__":
    main()
